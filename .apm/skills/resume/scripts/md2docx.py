#!/usr/bin/env python3
"""Convert a markdown resume to ATS-compliant, visually polished DOCX.

Usage: python3 md2docx.py <input.md> <output.docx>

Design: single column, no tables/textboxes/headers. Strong visual hierarchy
via font size, weight, color, spacing, and paragraph borders. Right-aligned
dates via tab stops. Metrics auto-bolded in bullets.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Calibri"
C_NAME = RGBColor(0x1B, 0x3A, 0x5C)
C_HEADING = RGBColor(0x1B, 0x3A, 0x5C)
C_BODY = RGBColor(0x33, 0x33, 0x33)
C_SUBTLE = RGBColor(0x55, 0x55, 0x55)
C_BULLET = RGBColor(0x3D, 0x6B, 0xA8)
C_METRIC = RGBColor(0x1B, 0x3A, 0x5C)
RULE_HEX = "1B3A5C"
PAGE_WIDTH = 6.6  # usable width in inches at 0.7" margins


def sp(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line


def border_bottom(p, color=RULE_HEX, sz="4"):
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    b = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): sz,
        qn("w:space"): "2", qn("w:color"): color,
    })
    pBdr.append(b)
    pPr.append(pBdr)


def r(p, text, size=10.5, bold=False, italic=False, color=C_BODY):
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return run


def doc_init():
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = FONT
    s.font.size = Pt(10.5)
    s.font.color.rgb = C_BODY
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    s.paragraph_format.line_spacing = 1.15
    for sec in doc.sections:
        sec.top_margin = Inches(0.55)
        sec.bottom_margin = Inches(0.45)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)
    return doc


def emit_name(doc, name):
    p = doc.add_paragraph()
    r(p, name, size=22, bold=True, color=C_NAME)
    sp(p, after=2)


def emit_contact(doc, line):
    p = doc.add_paragraph()
    parts = [x.strip() for x in line.split("|")] if "|" in line else [line]
    for i, part in enumerate(parts):
        r(p, part, size=9.5, color=C_SUBTLE)
        if i < len(parts) - 1:
            r(p, "   ·   ", size=9.5, color=C_BULLET)
    sp(p, after=1)


def emit_section(doc, text):
    p = doc.add_paragraph()
    sp(p, before=18, after=6)
    r(p, text.upper(), size=11, bold=True, color=C_HEADING)
    border_bottom(p)


def emit_job(doc, title, company, dates):
    p = doc.add_paragraph()
    sp(p, before=10, after=0)
    r(p, title, size=11, bold=True, color=C_NAME)
    if company:
        r(p, "  —  ", size=10, color=C_SUBTLE)
        r(p, company, size=10, color=C_BODY)
    if dates:
        from docx.enum.text import WD_TAB_ALIGNMENT
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(PAGE_WIDTH), WD_TAB_ALIGNMENT.RIGHT
        )
        r(p, "\t", size=10)
        r(p, dates, size=9.5, italic=True, color=C_SUBTLE)


METRIC_PAT = re.compile(
    r"(\$[\d,.]+[MKB+]*\+?|"
    r"\d+\+?\s*(?:years?|accounts?|customers?|engagements?|hackathons?|"
    r"students?|panels?|enterprises?|events?|verticals?|institutions?|"
    r"countries|certific|defens))"
)


def emit_bullet(doc, text):
    p = doc.add_paragraph()
    sp(p, before=1, after=2, line=1.12)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.18)
    r(p, "▸  ", size=8.5, color=C_BULLET)
    for i, part in enumerate(METRIC_PAT.split(text)):
        if i % 2 == 1:
            r(p, part, size=10.5, bold=True, color=C_METRIC)
        else:
            r(p, part, size=10.5, color=C_BODY)


def emit_competency(doc, text):
    p = doc.add_paragraph()
    sp(p, before=2, after=3, line=1.12)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.18)
    if ":" in text:
        cat, skills = text.split(":", 1)
        r(p, "▸  ", size=8.5, color=C_BULLET)
        r(p, cat.strip() + ":  ", size=10.5, bold=True, color=C_HEADING)
        r(p, skills.strip(), size=10.5, color=C_BODY)
    else:
        r(p, "▸  ", size=8.5, color=C_BULLET)
        r(p, text.strip(), size=10.5, color=C_BODY)


def emit_cert(doc, text):
    p = doc.add_paragraph()
    sp(p, before=1, after=2, line=1.1)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.18)
    r(p, "▸  ", size=8.5, color=C_BULLET)
    if " — " in text:
        cert, rest = text.split(" — ", 1)
        r(p, cert.strip(), size=10.5, bold=True, color=C_BODY)
        r(p, " — ", size=10.5, color=C_SUBTLE)
        r(p, rest.strip(), size=10.5, color=C_SUBTLE)
    else:
        r(p, text.strip(), size=10.5, color=C_BODY)


def emit_body(doc, text, size=10.5, color=C_BODY, indent=0, line_sp=1.18):
    p = doc.add_paragraph()
    sp(p, before=2, after=5, line=line_sp)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r(p, text, size=size, color=color)


def convert(md_path, docx_path):
    lines = Path(md_path).read_text(encoding="utf-8").split("\n")
    doc = doc_init()
    i = 0
    name_done = False
    section = ""

    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        # h1 = name
        if line.startswith("# ") and not name_done:
            emit_name(doc, line[2:].strip())
            name_done = True
            i += 1
            continue

        # contact lines before first h2
        if name_done and not line.startswith("#") and section == "":
            emit_contact(doc, line)
            i += 1
            continue

        # h2 = section
        if line.startswith("## "):
            heading = line[3:].strip()
            section = heading.lower()
            emit_section(doc, heading)
            i += 1
            continue

        # h3 = job title
        if line.startswith("### "):
            raw = line[4:].strip()
            parts = [x.strip() for x in raw.split("|")]
            title = parts[0] if parts else raw
            company = parts[1] if len(parts) > 1 else ""
            dates = parts[2] if len(parts) > 2 else ""
            emit_job(doc, title, company, dates)
            i += 1
            continue

        # bullets
        if line.startswith("- ") or line.startswith("* ") or line.startswith("→ "):
            txt = re.sub(r"^[-*→]\s+", "", line)
            if "competenc" in section:
                emit_competency(doc, txt)
            elif "certific" in section:
                emit_cert(doc, txt)
            elif "honor" in section or "award" in section:
                emit_cert(doc, txt)
            else:
                emit_bullet(doc, txt)
            i += 1
            continue

        # body text
        if not line.startswith("#"):
            emit_body(doc, line)
            i += 1
            continue

        i += 1

    doc.save(docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
